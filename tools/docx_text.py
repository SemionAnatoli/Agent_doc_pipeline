from __future__ import annotations

from pathlib import Path


def extract_docx_plain_text(path: Path, max_chars: int = 350_000) -> str:
    """Extract plain text from a .docx file using python-docx."""
    from docx import Document  # lazy import — optional dependency

    doc = Document(str(path))
    parts: list[str] = []
    total = 0
    for para in doc.paragraphs:
        t = para.text
        if not t:
            continue
        parts.append(t)
        total += len(t)
        if total >= max_chars:
            break
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    parts.append(t)
                    total += len(t)
                    if total >= max_chars:
                        break
    return "\n".join(parts)


def count_docx_pages(path: Path) -> int:
    """Approximate page count from paragraph count (roughly 40 paragraphs per page)."""
    try:
        from docx import Document
        doc = Document(str(path))
        paras = sum(1 for p in doc.paragraphs if p.text.strip())
        return max(1, paras // 40 + (1 if paras % 40 else 0))
    except Exception:
        return 1
