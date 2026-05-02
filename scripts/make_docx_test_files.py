"""
Генератор тестовых .docx файлов для папки TEST/.
Запуск: .venv\\Scripts\\python scripts/make_docx_test_files.py
"""
from __future__ import annotations
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "TEST"
OUT.mkdir(exist_ok=True)
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def row(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    run_l = p.add_run(f"{label}: ")
    run_l.bold = True
    p.add_run(value)


def separator(doc: Document) -> None:
    doc.add_paragraph("─" * 60)


# ── File 5: Invoice DOCX — CRM match (Sberbank) ─────────────────────────────
def make_invoice_docx() -> None:
    doc = Document()
    doc.add_heading("INVOICE No. 2025-05-042", 0)

    separator(doc)
    heading(doc, "Seller Details", 2)
    row(doc, "Company", "PJSC Sberbank")
    row(doc, "INN", "7702070139")
    row(doc, "KPP", "770201001")
    row(doc, "Address", "Moscow, ul. Vavilova 19")
    row(doc, "Bank", "Sberbank, BIK 044525225")

    separator(doc)
    heading(doc, "Buyer Details", 2)
    row(doc, "Company", "OOO Digital Solutions")
    row(doc, "INN", "7736050003")
    row(doc, "KPP", "773601001")
    row(doc, "Address", "Moscow, Kutuzovsky pr. 32")

    separator(doc)
    heading(doc, "Invoice Details", 2)
    row(doc, "Invoice date", "05.05.2025")
    row(doc, "Contract No.", "SB-2025/IT-019 dated 01.01.2025")
    row(doc, "Payment due", "20.05.2025")

    separator(doc)
    heading(doc, "Services", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Description"
    hdr[1].text = "Qty"
    hdr[2].text = "Amount (rub.)"
    items = [
        ("Cloud infrastructure services", "1 month", "280 000,00 rub."),
        ("API integration support", "40 hours", "120 000,00 rub."),
        ("Security audit", "1 project", "95 000,00 rub."),
        ("Technical documentation", "1 set", "45 000,00 rub."),
    ]
    for desc, qty, amt in items:
        r = table.add_row().cells
        r[0].text = desc
        r[1].text = qty
        r[2].text = amt

    separator(doc)
    doc.add_paragraph("Subtotal (excl. VAT):   540 000,00 rub.")
    doc.add_paragraph("VAT 20%:                108 000,00 rub.")
    p = doc.add_paragraph()
    run = p.add_run("TOTAL DUE:              648 000,00 rub.")
    run.bold = True

    separator(doc)
    doc.add_paragraph("Director: Gref G.O.  _______________")
    doc.add_paragraph("Chief Accountant: Ivanova E.S.  _______________")
    doc.add_paragraph("INN seller: 7702070139   INN buyer: 7736050003")

    path = OUT / "05_invoice_DOCX_sberbank_CRM_match.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


# ── File 6: Contract DOCX — unknown INN → review required ───────────────────
def make_contract_docx() -> None:
    doc = Document()
    doc.add_heading("SERVICE AGREEMENT No. 2025/AG-117", 0)

    separator(doc)
    row(doc, "Agreement date", "01.04.2025")
    row(doc, "Valid until", "31.12.2025")

    separator(doc)
    heading(doc, "Party A — Customer", 2)
    row(doc, "Company", "OOO NovaTech Innovations")
    row(doc, "INN", "3328012456")
    row(doc, "Address", "Vladimir, ul. Lenina 44")

    separator(doc)
    heading(doc, "Party B — Contractor", 2)
    row(doc, "Company", "IP Kolesnikov Andrey")
    row(doc, "INN", "332800987654")
    row(doc, "Address", "Vladimir, pr. Stroiteley 12, apt. 5")

    separator(doc)
    heading(doc, "Subject of Agreement", 2)
    doc.add_paragraph(
        "Party B agrees to provide software development services for Party A, "
        "including design, implementation and testing of a mobile application "
        "for logistics management."
    )

    separator(doc)
    heading(doc, "Financial Terms", 2)
    row(doc, "Total contract value", "1 200 000,00 rub.")
    row(doc, "Advance payment (30%)", "360 000,00 rub.")
    row(doc, "Payment on milestone 1", "480 000,00 rub.")
    row(doc, "Final payment", "360 000,00 rub.")
    row(doc, "Payment schedule", "15.05.2025, 15.08.2025, 15.11.2025")

    separator(doc)
    doc.add_paragraph("Signed:")
    doc.add_paragraph("Party A: _______________   01.04.2025")
    doc.add_paragraph("Party B: _______________   01.04.2025")
    doc.add_paragraph("INN Party A: 3328012456   INN Party B: 332800987654")

    path = OUT / "06_contract_DOCX_unknown_INN_review_required.docx"
    doc.save(str(path))
    print(f"  Created: {path.name}")


if __name__ == "__main__":
    make_invoice_docx()
    make_contract_docx()
    print(f"\nDocx test files saved to: {OUT}")
    print("\nExpected pipeline results:")
    print("  05_invoice_DOCX_sberbank_CRM_match.docx   -> All 5 steps GREEN")
    print("  06_contract_DOCX_unknown_INN_review.docx  -> Route: Review required")
